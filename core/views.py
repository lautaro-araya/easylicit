from django.shortcuts import render, redirect
from django.contrib.auth import login, logout, authenticate
from django.contrib.auth.forms import AuthenticationForm
import os
from django.contrib import messages
from django.http import HttpResponse
from django.core.files.storage import FileSystemStorage
from openpyxl import load_workbook, Workbook
from docx import Document
from xlrd import open_workbook
from .forms import RegistroClienteForm
from django.contrib.auth.decorators import login_required
from docx2pdf import convert


UPLOAD_DIR = 'uploads/'

def cerrar_sesion(request):
    logout(request)
    return redirect('iniciar_sesion')

def registro_cliente(request):
    if request.method == 'POST':
        form = RegistroClienteForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'El usuario ha sido registrado exitosamente.')
            return redirect('iniciar_sesion')
    else:
        form = RegistroClienteForm()
    
    return render(request, 'core/register.html', {'form': form})

@login_required
def home(request):
    return render(request, "core/home.html")

@login_required
def convert_excel(request):
    if request.method == 'POST' and request.FILES.get('xls_file'):
        # Guardar el archivo .xls subido
        xls_file = request.FILES['xls_file']
        if not os.path.exists(UPLOAD_DIR):
            os.makedirs(UPLOAD_DIR)  # Crear la carpeta si no existe
        filepath = os.path.join(UPLOAD_DIR, xls_file.name)
        with open(filepath, 'wb') as f:
            for chunk in xls_file.chunks():
                f.write(chunk)

        try:
            # Leer el archivo .xls usando xlrd
            workbook_xls = open_workbook(filepath)
            sheet_xls = workbook_xls.sheet_by_index(0)  # Selecciona la primera hoja

            # Crear un archivo .xlsx con openpyxl
            workbook_xlsx = Workbook()
            sheet_xlsx = workbook_xlsx.active

            # Copiar datos del archivo .xls al .xlsx
            for row in range(sheet_xls.nrows):
                row_data = [sheet_xls.cell_value(row, col) for col in range(sheet_xls.ncols)]
                sheet_xlsx.append(row_data)

            # Generar el archivo convertido en memoria
            response = HttpResponse(
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename="archivo_convertido.xlsx"'
            workbook_xlsx.save(response)
            return response

        except Exception as e:
            return HttpResponse(f"Error al convertir el archivo: {str(e)}", status=500)

    return render(request, "core/convert_excel.html")

def iniciar_sesion(request):
    errors = {}
    if request.method == "POST":
        form = AuthenticationForm(request, data=request.POST)
        if form.is_valid():
            nombre_usuario = form.cleaned_data.get("username")
            contraseña = form.cleaned_data.get("password")
            user = authenticate(username=nombre_usuario, password=contraseña)
            if user is not None:
                login(request, user)
                return redirect('home')
            else:
                errors['autenticación'] = "*Usuario o contraseña incorrectos."
        else:
            errors['autenticación'] = "Usuario o contraseña incorrectos."

    form = AuthenticationForm()

    return render(request, "core/login.html", {
        "form": form,
        "errors": errors,
    })

@login_required
def upload_excel(request):
    if request.method == 'POST' and request.FILES.get('excel_file'):
        # Guardar el archivo Excel en el servidor
        excel_file = request.FILES['excel_file']
        fs = FileSystemStorage(location=UPLOAD_DIR)
        filename = fs.save(excel_file.name, excel_file)
        filepath = os.path.join(UPLOAD_DIR, filename)

        # Leer la fila seleccionada
        selected_row = request.POST.get('selected_row')
        try:
            selected_row = int(selected_row)  # Asegúrate de que es un número válido
        except ValueError:
            return HttpResponse("Número de fila no válido. Por favor, inténtalo de nuevo.")
        
        workbook = load_workbook(filepath)
        sheet = workbook.active

        # Validar que la fila seleccionada sea válida
        try:
            row_data = sheet[selected_row]
        except IndexError:
            return HttpResponse("Número de fila fuera de rango. Por favor, verifica el archivo Excel.")

        # Extraer datos de la fila
        id_value = row_data[0].value
        nombre_value = row_data[1].value
        institucion_value = row_data[7].value
        presupuesto_value = row_data[8].value

        # Extraer datos adicionales del formulario
        datos_tecnicos = request.POST.get('tecnic')
        fecha_entrega = request.POST.get('fecha')
        plazo_entrega = request.POST.get('plazo')
        garantia = request.POST.get('garantia')
        subtotal = request.POST.get('subtotal')

        total = int(int(subtotal)*1.19)
        iva = int(total) - int(subtotal)

        # Cargar la plantilla Word
        template_path = 'core/plantilla.docx'  # Cambia esto a la ruta real
        doc = Document(template_path)

        # Reemplazar los marcadores en el cuerpo del documento
        for paragraph in doc.paragraphs:
            if '{id}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{id}', str(id_value))
            if '{institucion}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{institucion}', str(institucion_value))
            if '{presupuesto}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{presupuesto}', f"{presupuesto_value}")
            if '{nombre}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{nombre}', f"{nombre_value}")
            if '{datos_tecnicos}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{datos_tecnicos}', datos_tecnicos)
            if '{fecha}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{fecha}', fecha_entrega)
            if '{plazo}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{plazo}', f"{plazo_entrega}")
            if '{garantia}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{garantia}', f"{garantia}")
            if '{subtotal}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{subtotal}', f"{subtotal}")
            if '{iva}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{iva}', f"{iva}")
            if '{total}' in paragraph.text:
                paragraph.text = paragraph.text.replace('{total}', f"{total}")


        # Reemplazar los marcadores en encabezados y pies de página
        for section in doc.sections:
            # Encabezado
            header = section.header
            for paragraph in header.paragraphs:
                if '{id}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{id}', str(id_value))
                if '{institucion}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{institucion}', str(institucion_value))
                if '{presupuesto}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{presupuesto}', f"{presupuesto_value}")
                if '{nombre}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{nombre}', f"{nombre_value}")
            
            # Pie de página
            footer = section.footer
            for paragraph in footer.paragraphs:
                if '{id}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{id}', str(id_value))
                if '{institucion}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{institucion}', str(institucion_value))
                if '{presupuesto}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{presupuesto}', f"{presupuesto_value}")
                if '{nombre}' in paragraph.text:
                    paragraph.text = paragraph.text.replace('{nombre}', f"{nombre_value}")

        # Ruta para guardar el archivo DOCX generado
        generated_path = os.path.join(UPLOAD_DIR, 'documento_generado.docx')
        doc.save(generated_path)  # Guardar el archivo generado (ya lo tienes configurado previamente)

        # Ruta para el archivo PDF generado
        pdf_generated_path = os.path.join(UPLOAD_DIR, 'documento_generado.pdf')

        # Convertir el archivo DOCX a PDF
        convert(generated_path, pdf_generated_path)  # Convierte el DOCX a PDF

        # Verificar si el archivo PDF fue generado correctamente
        if not os.path.exists(pdf_generated_path):
            return HttpResponse("El archivo PDF no se generó correctamente.", status=500)
        
        with open(pdf_generated_path, 'rb') as f:
            response = HttpResponse(f.read(), content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename=cotizacion.pdf'
            return response
        # Redirigir a la misma página      


    return render(request, 'core/upload_excel.html')